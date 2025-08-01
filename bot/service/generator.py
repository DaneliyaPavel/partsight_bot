from __future__ import annotations

import datetime
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.text import WD_BREAK
from docxcompose.composer import Composer
from docxtpl import DocxTemplate

from bot.config import settings

# ─────────────────────── Шаблоны ────────────────────────
_TPL_FIRST = settings.TEMPLATE_PATH  # со шапкой
_TPL_TABLE = settings.TEMPLATE_TABLE_PATH  # только таблица

_OUT = Path("output")
_OUT.mkdir(exist_ok=True)


# ─────────────────────── Вспом. рендер ───────────────────
def _render_one(tpl_path: Path, ctx: Dict[str, str]) -> Path:
    """
    Рендерит один документ по шаблону, возвращает путь к tmp-файлу.
    """
    tpl = DocxTemplate(tpl_path)
    tpl.render(ctx)
    tmp = _OUT / f"_tmp_{ctx['material']}.docx"
    tpl.save(tmp)
    return tmp


# ─────────────────────── Основная сборка ─────────────────
def build_act(packs: List[Dict[str, str]]) -> Path:
    """
    • Первый лист — шаблон с заголовком.
    • Далее — только таблица. Каждая таблица с новой страницы.
    """
    first: bool = True
    composer: Composer | None = None
    master_doc: Document | None = None

    for p in packs:
        ctx = dict(
            material=p["material"],
            naimenovanie=p["naimenovanie"],
            obosnovanost="обоснованно",
            nalichie="отсутствие",
            vozmoznost=p["vozmoznost"],
            rekomendacii=p["rekomendacii"],
        )

        tpl_file = _TPL_FIRST if first else _TPL_TABLE
        tmp = _render_one(tpl_file, ctx)

        if first:
            master_doc = Document(tmp)
            composer = Composer(master_doc)
            first = False
        else:
            assert composer is not None
            assert master_doc is not None
            # добавляем разрыв страницы
            master_doc.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
            composer.append(Document(tmp))

        tmp.unlink()

    # финальная сохранёнка
    assert composer is not None
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    out = _OUT / f"Акт_{ts}.docx"
    composer.save(out)
    return out
